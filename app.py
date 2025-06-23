import streamlit as st
from agents.chain_agent import ContentChainAgent
import os
from dotenv import load_dotenv
import re
import io
import torch
import base64
from docx import Document
from docx.shared import Inches

def main():
    """
    ContentCrafter AI — Final Streamlit App

    ✅ Supports: [Insert image here: ...], (Image: ...), (Infographic: ...)
    ✅ Inline image rendering in UI and DOCX
    ✅ Uses RTX 4060 via torch.float16 if available
    ✅ Handles image fallback, CUDA OOM, unsupported formats (GIFs/videos)
    """

    load_dotenv()
    hf_token = os.getenv("HUGGINGFACE_API_TOKEN")

    generate_image = lambda prompt: "<!-- Image disabled -->"
    try:
        from diffusers import StableDiffusionPipeline
        from transformers import logging as hf_logging
        from PIL import Image
        from huggingface_hub import login

        hf_logging.set_verbosity_error()
        login(token=hf_token, add_to_git_credential=False)

        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        pipe = StableDiffusionPipeline.from_pretrained(
            "runwayml/stable-diffusion-v1-5",
            torch_dtype=torch.float16 if device.type == "cuda" else torch.float32,
            safety_checker=None,
            use_safetensors=True,
            low_cpu_mem_usage=False
        ).to(device)

        print(f"✅ Stable Diffusion ready on {device.type.upper()}")

        def generate_image(prompt):
            try:
                prompt = prompt.strip()[:150]
                with torch.no_grad():
                    image = pipe(prompt, num_inference_steps=30, guidance_scale=7.5).images[0]
                buf = io.BytesIO()
                image.save(buf, format="PNG")
                return f"data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}"
            except torch.cuda.OutOfMemoryError:
                pipe.to("cpu")
                with torch.no_grad():
                    image = pipe(prompt, num_inference_steps=20, guidance_scale=7.5).images[0]
                buf = io.BytesIO()
                image.save(buf, format="PNG")
                return f"data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}"
            except Exception as e:
                print(f"[Image Error] {prompt}: {e}")
                return f"<!-- Image failed: {e} -->"

    except Exception as e:
        print(f"[Image pipeline fallback] {e}")

    st.set_page_config(page_title="ContentCrafter AI", page_icon="🧠", layout="wide")
    st.title("🧠 ContentCrafter AI")
    st.subheader("Generate intelligent blog content + images")

    topic_input = st.text_area(
        "📝 Enter one or more topics (comma or line-separated):",
        placeholder="e.g., AI in Healthcare, Space Farming",
        height=100
    )

    if st.button("🚀 Generate Content"):
        if not topic_input.strip():
            st.warning("Please enter at least one topic.")
            return

        topics = [t.strip() for t in topic_input.replace("\n", ",").split(",") if t.strip()]
        agent = ContentChainAgent(model_name="gemini-1.5-flash")

        with st.spinner("Generating content..."):
            results = agent.run_chain(",".join(topics))

        for topic, output in results.items():
            st.markdown(f"## 🧠 Topic: `{topic}`")
            if "error" in output:
                st.error(output["error"])
                continue

            blog_title = output.get("blog_title", topic)
            drafts = {
                "Original Draft": output.get("initial_draft", ""),
                "Editor Feedback": output.get("feedback", ""),
                "Improved Draft": output.get("blog_post", ""),
                "Structural Feedback": output.get("structural_feedback", ""),
                "Second Draft": output.get("second_draft", ""),
                "Final Post": output.get("edited_post", "")
            }

            st.markdown("### 📌 Content Plan")
            st.code(output.get("plan", ""), language="markdown")

            def render_with_images(text, label):
                def embed(prompt):
                    if "video" in prompt.lower() or "gif" in prompt.lower():
                        st.warning(f"⚠️ Skipping video/GIF: {prompt}")
                        return f"⚠️ Video/GIF not supported: {prompt}"
                    img_url = generate_image(prompt)
                    return f'<img src="{img_url}" alt="{prompt}" width="300">' if img_url.startswith("data:image") else f"⚠️ Failed: {prompt}"

                html = re.sub(r"\[Insert image here: (.*?)\]", lambda m: embed(m.group(1)), text)
                html = re.sub(r"\(Infographic: (.*?)\)", lambda m: embed(m.group(1)), html)
                html = re.sub(r"\(Image: (.*?)\)", lambda m: embed(m.group(1)), html)
                html = re.sub(r"\[Insert short video or animated GIF here: (.*?)\]", lambda m: embed(m.group(1)), html)
                html = re.sub(r"\s*\n", " ", html)
                st.markdown(f"### {label}")
                st.markdown(html, unsafe_allow_html=True)
                return html

            for label, content in drafts.items():
                if label in ["Original Draft", "Editor Feedback", "Structural Feedback"]:
                    st.markdown(f"### {label}")
                    st.code(content, language="markdown")
                else:
                    drafts[label] = render_with_images(content, label)

            try:
                doc = Document()
                doc.add_heading(blog_title, 0)
                final_html = drafts["Final Post"]
                for line in final_html.split("<img"):
                    if "src=" in line:
                        try:
                            b64 = re.search(r'src="data:image/png;base64,(.*?)"', line).group(1)
                            img_data = base64.b64decode(b64)
                            with io.BytesIO() as img_io:
                                img_io.write(img_data)
                                img_io.seek(0)
                                doc.add_picture(img_io, width=Inches(4))
                        except Exception as e:
                            doc.add_paragraph(f"[Image failed: {e}]")
                    else:
                        clean = re.sub(r"<.*?>", "", line).strip()
                        if clean:
                            doc.add_paragraph(clean)

                out_io = io.BytesIO()
                doc.save(out_io)
                out_io.seek(0)
                st.download_button(
                    label=f"💾 Download DOCX for '{topic}'",
                    data=out_io.read(),
                    file_name=f"{blog_title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"⚠️ DOCX Export Failed: {e}")

    st.markdown("---")
    st.caption("Built for the Google Cloud Agent Hackathon")

if __name__ == "__main__":
    main()
